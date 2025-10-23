from docx import Document
import os

def convert_txt_to_docx(txt_file_path, docx_file_path):
    # Read the text file
    with open(txt_file_path, 'r', encoding='utf-8') as txt_file:
        text_content = txt_file.read()

    # Create a new Word document
    doc = Document()

    # Split the text into paragraphs (assuming paragraphs are separated by double newlines)
    paragraphs = text_content.split('\n\n')

    for paragraph in paragraphs:
        # Handle different types of content
        if paragraph.strip().startswith('CHAPTER') or paragraph.strip().startswith('FIGURE') or paragraph.strip().startswith('TABLE'):
            # Add as heading
            doc.add_heading(paragraph.strip(), level=1)
        elif paragraph.strip().startswith('LIST OF') or paragraph.strip().startswith('ABSTRACT') or paragraph.strip().startswith('DECLARATION') or paragraph.strip().startswith('APPROVAL') or paragraph.strip().startswith('ACKNOWLEDGEMENTS') or paragraph.strip().startswith('DEDICATION') or paragraph.strip().startswith('REFERENCES') or paragraph.strip().startswith('APPENDICES'):
            doc.add_heading(paragraph.strip(), level=1)
        elif paragraph.strip().startswith('Figure ') or paragraph.strip().startswith('Table '):
            # Add as caption
            doc.add_paragraph(paragraph.strip(), style='Caption')
        elif paragraph.strip().startswith('- ') or paragraph.strip().startswith('• '):
            # Bullet points
            doc.add_paragraph(paragraph.strip(), style='List Bullet')
        elif '|' in paragraph and len(paragraph.split('|')) > 3:
            # Table-like content
            lines = paragraph.strip().split('\n')
            if len(lines) > 1:
                # Create table
                num_cols = len(lines[0].split('|')) - 1  # Subtract 1 for empty first/last
                if num_cols > 1:
                    table = doc.add_table(rows=len(lines), cols=num_cols)
                    table.style = 'Table Grid'
                    for i, line in enumerate(lines):
                        cells = line.split('|')[1:-1]  # Skip first and last empty
                        for j, cell in enumerate(cells):
                            if j < num_cols:
                                table.cell(i, j).text = cell.strip()
                else:
                    doc.add_paragraph(paragraph.strip())
            else:
                doc.add_paragraph(paragraph.strip())
        else:
            # Regular paragraph
            doc.add_paragraph(paragraph.strip())

    # Save the document
    doc.save(docx_file_path)
    print(f"Converted {txt_file_path} to {docx_file_path}")

if __name__ == "__main__":
    txt_file = "Documentation_final_updated.txt"
    docx_file = "Documentation_final_updated.docx"

    if os.path.exists(txt_file):
        convert_txt_to_docx(txt_file, docx_file)
    else:
        print(f"Error: {txt_file} not found")